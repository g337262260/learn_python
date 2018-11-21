#! python3
""""
    #:处理word文档
    #：https://python-docx.readthedocs.io/en/latest/user/quickstart.html
"""

import docx


def readWord():
    doc = docx.Document('panda.docx')
    # 获取word的段落数量
    graphs = len(doc.paragraphs)
    # 打印每一段文字的内容，表格不支持
    for i in range(0,graphs):
        print('text:'+doc.paragraphs[i].text)


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in  doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def writeDocx(filename,title,size=3):
    doc = docx.Document()
    # 添加标题
    doc.add_heading(title,size)
    print('Start Write')
    doc.add_paragraph('Hello world')
    # 换行
    # doc.paragraphs[0].runs[0].add_break()
    # 换页
    doc.add_page_break()
    doc.add_paragraph('How are you ?')
    doc.save(filename)
    print('Write Done')



if __name__ == '__main__':
    writeDocx('text.docx','Hello world!',size=0)